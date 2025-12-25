from pathlib import Path
from typing import List
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader
import docx
from app.config import settings


def load_pdf(path: Path) -> List[Document]:
    """下面的代码将pdf分割成单独的页面，每一个页面的文本被封装成一个Document放入list"""
    reader = PdfReader(str(path))
    docs = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            docs.append(Document(
                page_content=text,
                metadata={"source": str(path), "page": i+1}
            ))
    return docs


def load_docx(path: Path) -> List[Document]:
    d = docx.Document(str(path))
    text = "\n".join(p.text for p in d.paragraphs if p.text.strip())
    return [Document(page_content=text, metadata={"source": str(path)})] if text else []


def load_txt(path: Path) -> List[Document]:
    """加载文本文件，保留原始格式"""
    text = path.read_text(encoding="utf-8")
    if not text.strip():  # 空文件
        return []
    return [Document(page_content=text, metadata={"source": str(path)})]


def load_docs(dir_path: str) -> List[Document]:
    p = Path(dir_path)
    docs: List[Document] = []
    for f in p.rglob("*"):
        if f.suffix.lower() == ".pdf":
            docs.extend(load_pdf(f))
        elif f.suffix.lower() in [".docx", ".doc"]:
            docs.extend(load_docx(f))
        elif f.suffix.lower() in [".md", ".txt"]:
            docs.append(Document(page_content=f.read_text(encoding="utf-8"),
                                 metadata={"source": str(f)}))
    return docs

def split_docs(docs: List[Document]) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap
    )
    return splitter.split_documents(docs)

# if __name__ == "__main__":
#     for l in load_docs('../../data/docs'):
#         print('-------------begin--------------')
#         print(l)
#         print('--------------end-------------')


# if __name__ == "__main__":
#     print(len(split_docs(load_docs('../../data'))))

# 第一个函数主要是为了应对后续文件单独追加而设计的，他就是处理一个单独的文件而已。
def load_single_file(path: Path) -> List[Document]:
    """根据文件后缀加载文件，返回LangChain的 Document列表"""
    suf = path.suffix.lower()
    if suf == ".pdf":
        return load_pdf(path)
    if suf in [".docx", ".doc"]:
        return load_docx(path)
    if suf in [".md", ".txt"]:
        text = path.read_text(encoding="utf-8")
        return [Document(page_content=text, metadata={"source": str(path)})] if text.strip() else []
    return []

# 第二个函数主要是把一批Document切成小块，并且给每一小块贴上权限标签visibility和文档ID。
def split_with_visibility(docs: List[Document], visibility: str, doc_id: str | None = None) -> List[Document]:
    chunks = split_docs(docs)
    for c in chunks:
        c.metadata = dict(c.metadata or {})
        c.metadata["visibility"] = visibility
        if doc_id:
            c.metadata["doc_id"] = doc_id
    return chunks


